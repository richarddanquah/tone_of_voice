from langchain.chains import LLMChain
from langchain.llms import OpenAI
from langchain.prompts import PromptTemplate
from langchain.vectorstores import Chroma
from langchain_openai import OpenAIEmbeddings
from langchain.document_loaders import TextLoader
from langchain.text_splitter import RecursiveCharacterTextSplitter
from dotenv import load_dotenv
import os
from docx import Document
from typing import List


load_dotenv()
os.environ["OPENAI_API_KEY"] =os.getenv("OPENAI_API_KEY")


# Prompts
signature_prompt = PromptTemplate.from_template("""
Analyze the tone of the following text and extract:
- Tone
- Language Style
- Level of Formality
- Forms of Address
- Emotional Appeal

Text:
{text}
""")

rewrite_prompt = PromptTemplate.from_template("""
Rewrite this text to reflect the following tone signature:
{signature}

Text:
{text}
""")

evaluation_prompt = PromptTemplate.from_template("""
Evaluate how well the rewritten text aligns with the tone signature below. Return:
- Fluency
- Authenticity
- Tone Alignment
- Readability
- Score (0 to 1)
- Strengths
- Suggestions

Signature:
{signature}

Original:
{original_text}

Rewritten:
{rewritten_text}
""")

#analyze tone
def analyze_tone(text:str): 
    chain = LLMChain(llm=OpenAI(temperature=0), prompt=signature_prompt)
    return chain.run({"text": text})

#rewrite_with_signature
def rewrite_with_signature(text: str, signature: str):
    chain = LLMChain(llm=OpenAI(temperature=0.7), prompt=rewrite_prompt)
    return chain.run({"text": text, "signature": signature})

#evaluate_tone
def evaluate_tone(original: str, rewritten: str, signature: str):
    chain = LLMChain(llm=OpenAI(temperature=0), prompt=evaluation_prompt)
    return chain.run({"original_text": original, "rewritten_text": rewritten, "signature": signature})

class WordDocumentLoader:
    """Custom loader for Word documents"""
    def __init__(self, file_path: str):
        self.file_path = file_path

    def load(self) -> List[str]:
        """Load and extract text from a Word document"""
        doc = Document(self.file_path)
        texts = []
        
        # Extract text from paragraphs
        for para in doc.paragraphs:
            if para.text.strip():  # Only add non-empty paragraphs
                texts.append(para.text)
        
        # Extract text from tables
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if cell.text.strip():  # Only add non-empty cells
                        texts.append(cell.text)
        
        return texts

def process_word_document(file_path: str):
    """Process a Word document and store it in the vector database for brand voice analysis."""
    # Load the document
    loader = WordDocumentLoader(file_path)
    texts = loader.load()
    
    # Split the text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    documents = text_splitter.create_documents(texts)
    
    # Create embeddings and store in Chroma
    embeddings = OpenAIEmbeddings()
    vectorstore = Chroma.from_documents(
        documents=documents,
        embedding=embeddings,
        persist_directory="./chroma_db"
    )
    
    return vectorstore

def process_document(file_path: str):
    """Process a document and store it in the vector database for brand voice analysis."""
    # Load the document
    loader = TextLoader(file_path)
    documents = loader.load()
    
    # Split the text into chunks
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=1000,
        chunk_overlap=200
    )
    texts = text_splitter.split_documents(documents)
    
    # Create embeddings and store in Chroma
    embeddings = OpenAIEmbeddings()
    vectorstore = Chroma.from_documents(
        documents=texts,
        embedding=embeddings,
        persist_directory="./chroma_db"
    )
    
    return vectorstore

def find_similar_texts(text: str, vectorstore, k: int = 3):
    """Find similar texts from the vector database that match the brand voice."""
    embeddings = OpenAIEmbeddings()
    similar_texts = vectorstore.similarity_search(text, k=k)
    return similar_texts

def analyze_brand_voice(corpus_paths: list[str]):
    """Analyze multiple documents to build a comprehensive brand voice profile."""
    all_texts = []
    
    for path in corpus_paths:
        vectorstore = process_document(path)
        # Get all documents from the vectorstore
        texts = vectorstore.get()
        all_texts.extend(texts)
    
    # Combine all texts for analysis
    combined_text = " ".join([doc.page_content for doc in all_texts])
    
    # Analyze the combined text to get a comprehensive brand voice signature
    return analyze_tone(combined_text)

def analyze_press_releases(docx_path: str = "Press Releases Examples.docx"):
    """
    Process the Press Releases Examples.docx file, analyze its tone, and return the signature.
    """
    vectorstore = process_word_document(docx_path)
    documents = vectorstore.get()
    # The documents are already strings, so we can join them directly
    combined_text = " ".join(documents)
    signature = analyze_tone(combined_text)
    return signature